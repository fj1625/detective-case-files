# -*- coding: utf-8 -*-
"""
Generate the answer key Word doc for the detective case files.

Each case lists the correct answer for every interactive page, so the web
pages can be verified quickly. To add a new case (or page), append to CASES
below and re-run:  python answer_key_generator.py

Output: 答案验证表 Answer Key.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "答案验证表 Answer Key.docx"

# ---------------------------------------------------------------------------
# Answer data — one dict per case. Add a new case dict to extend.
# timeline: the 5 events already in correct chronological order (1->5).
# suspects: suspect name -> list of its clue texts (the folder it belongs in).
# evidence: real = tiles to DROP on the board; distractor = tile to REJECT.
# culprit: the suspect to drag into the accusation zone.
# ---------------------------------------------------------------------------
CASES = [
    {
        "num": "01",
        "title": "The Lost Picture",
        "title_cn": "丢失的照片",
        "cover": {
            "Client": "Annie",
            "Missing": "A picture of Fang",
            "Location": "Down the street",
        },
        "timeline": [
            "Annie called Felix. Her Fang picture was gone.",
            "Felix checked the yellow room. No picture.",
            "Fang buried bones but stayed in the yard.",
            "Felix found Rosamond's lost cat.",
            "Harry's orange monster hid the picture.",
        ],
        "suspects": {
            "Annie": ["lost a yellow picture of her dog", "has a yellow room"],
            "Fang": ["big dog with big teeth", "never leaves the yard without a leash"],
            "Rosamond": ["has four cats named Hex", "only likes cats"],
            "Harry": ["paints everything red", "painted an orange monster with three heads"],
        },
        "evidence_real": ["🎨 Wet Yellow Paint", "🖌️ Red Paintbrush", "👹 Monster Drawing"],
        "evidence_distractor": "🥞 Pancakes",
        "deduction": "Wet yellow paint + red paint = orange monster; the monster had dog ears and a tail = Fang.",
        "culprit": "Harry",
    },
    {
        "num": "02",
        "title": "The Garbage Snatcher",
        "title_cn": "翻垃圾的家伙",
        "cover": {
            "Client": "Oliver",
            "Missing": "Garbage, snatched every night",
            "Location": "Oliver's garbage can",
        },
        "timeline": [
            "Oliver said a snatcher tipped his can each night.",
            "Felix asked Rosamond. She hates garbage.",
            "Esmeralda said no one would go near Oliver.",
            "Felix hid in the garbage can at night.",
            "Sludge knocked off the cover. Sludge did it!",
        ],
        "suspects": {
            "Rosamond": ["would eat 2000 things before garbage", "her cats eat cat food, not garbage"],
            "Esmeralda": ["keeps her mouth open to say wise things", "would never go near Oliver the pest"],
            "Skunk": ["found near garbage in a field", "sprayed Felix with a big stink"],
            "Sludge": ["sniffed the garbage can a lot", "was tired of pancakes and hungry"],
        },
        "evidence_real": ["🚷 No Person", "🌙 Night Animal", "👃 Sludge Sniffs"],
        "evidence_distractor": "🦨 Smelly Skunk (ruled-out lead)",
        "deduction": "No person would go near Oliver, so a night animal did it; Sludge sniffed and knocked the cover off = Sludge.",
        "culprit": "Sludge",
    },
    {
        "num": "03",
        "title": "The Lost List",
        "title_cn": "丢失的清单",
        "cover": {
            "Client": "Claude",
            "Missing": "A grocery list",
            "Location": "Between Claude's house and the store",
        },
        "timeline": [
            "Claude lost his grocery list before lunch.",
            "Felix drew a map and walked every street.",
            "The wind blew the map toward Rosamond's house.",
            "Felix made pancakes — the list was all baking stuff.",
            "Rosamond's \"cat pancake recipe\" was the lost list!",
        ],
        "suspects": {
            "Claude": ["always loses things — even his way", "his dad wrote the grocery list"],
            "Fang": ["held a paper tight in his teeth", "barked back, then dropped the paper"],
            "Sludge": ["sniffed up and down every street", "barked funny to help Felix"],
            "Rosamond": ["was all covered in white flour", "found a new cat pancake recipe today"],
        },
        "evidence_real": ["🥞 Pancake Stuff", "🐟 Tuna Fish", "🌬️ The Wind"],
        "evidence_distractor": "🗺️ Fang's Paper (just Felix's map)",
        "deduction": "The list was all pancake stuff + tuna fish (cats love tuna) = looks like a cat pancake recipe; the wind blew it to Rosamond, who kept it.",
        "culprit": "Rosamond",
    },
    {
        "num": "04",
        "title": "The Phony Clue",
        "title_cn": "假线索",
        "cover": {
            "Client": "Felix himself",
            "Missing": "Who sent the torn invitation?",
            "Location": "Felix's doorstep & around the block",
        },
        "timeline": [
            "A torn scrap saying \"VITA\" was on Felix's doorstep.",
            "A pancake lured Big Hex down; Felix grabbed a tree piece.",
            "The pieces joined to read: INVITATION — COME AT THREE.",
            "Finley dropped a \"phony clue\" into the sewer.",
            "Wet ink matched — Finley wrote it! Felix went at three.",
        ],
        "suspects": {
            "Rosamond": ["owns four cats with sixteen sharp claws", "said Big Hex sat in his tree all morning"],
            "Annie": ["was out walking her dog Fang", "always said Felix is a great detective"],
            "Pip": ["hardly ever says a single word", "bet that Felix WOULD solve it by three"],
            "Finley": ["said, \"Maybe he's great, maybe he's not.\"", "dropped a paper into the sewer and walked away"],
        },
        "evidence_real": ["📄 \"VITA\" Scrap", "💧 Wet Paper", "🅴 The Funny \"E\""],
        "evidence_distractor": "🐱 Big Hex the cat (ruled-out lead — tears paper, can't write)",
        "deduction": "\"VITA\" is part of INVITATION; wet thin paper lets the ink show through; the funny \"E\" matches Finley's writing = Finley wrote it.",
        "culprit": "Finley",
    },
    {
        "num": "05",
        "title": "The Sticky Case",
        "title_cn": "黏糊糊的案子",
        "cover": {
            "Client": "Claude",
            "Missing": "A tiny stegosaurus stamp",
            "Location": "Claude's table → Rosamond's yard sale",
        },
        "timeline": [
            "Claude's tiny stegosaurus stamp vanished from his table.",
            "After the rain, Pip came over and stepped in puddles.",
            "The wet shoe pressed the stamp's sticky side — it stuck!",
            "Pip swapped his wet shoes at Rosamond's Swap Table.",
            "Annie bought the shoe — Fang was chewing the stamp!",
        ],
        "suspects": {
            "Annie": ["bought the wet shoes for Fang to chew", "said the stamp looks like Fang's smile"],
            "Pip": ["came over after the rain had stopped", "swapped his wet shoes for dry slippers"],
            "Rosamond": ["was holding a yard sale with a Swap Table", "sold Pip's shoes to Annie for ten cents"],
            "Fang": ["was chewing a shoe in Annie's yard", "the lost stamp was stuck to the shoe's sole"],
        },
        "evidence_real": ["💧 Rain & Puddles", "👟 Pip's Wet Shoe", "🦕 Sticky Stamp"],
        "evidence_distractor": "🥫 Empty Tuna Can (silly yard-sale prop)",
        "deduction": "Rain → puddles → Pip's wet shoe pressed the stamp's sticky side, so it stuck to his shoe and rode off (then swapped, sold, and chewed by Fang). Culprit = Pip, by accident.",
        "culprit": "Pip",
    },
    {
        "num": "06",
        "title": "The Missing Key",
        "title_cn": "丢失的钥匙",
        "cover": {
            "Client": "Annie",
            "Missing": "Her shiny silver house key",
            "Location": "Annie's house — Fang's birthday party",
        },
        "timeline": [
            "Annie left her key on a table and went to buy Fang's surprise.",
            "Rosamond and her four cats were left alone in the house.",
            "Annie returned — house locked, Rosamond gone, key missing.",
            "Felix hunted shiny places: Oliver's, a bank, a garbage can.",
            "The key hid among shiny silver — on Fang's collar!",
        ],
        "suspects": {
            "Rosamond": ["was left alone in Annie's house", "left a strange riddle poem on the door"],
            "Oliver": ["collects one new shiny thing each week", "this week it was shiny eels, not keys"],
            "Fang": ["wore a new collar full of shiny silver", "stayed safe in the yard all day"],
            "Sludge": ["sniffed every garbage can", "helped Felix pry off the can's cover"],
        },
        "evidence_real": ["🔑 Silver Key", "✨ Other Shiny Silver", "🐕 Fang's Round Collar"],
        "evidence_distractor": "🏦 The Bank (ruled-out lead — round & shiny, but too normal for Rosamond)",
        "deduction": "A shiny silver key is easy to hide among OTHER shiny silver things; Fang's collar is round, big, and safe with silver tags, so Rosamond hung the key there — inches from Fang's teeth.",
        "culprit": "Rosamond",
    },
    {
        "num": "07",
        "title": "The Snowy Trail",
        "title_cn": "雪地足迹",
        "cover": {
            "Client": "Rosamond",
            "Missing": "Felix's birthday present (she won't say what!)",
            "Location": "The snowy trail to Felix's house",
        },
        "timeline": [
            "Rosamond pulled Felix's present and 4 cats on her sled.",
            "Under a tree, the sled felt lighter — the present was gone.",
            "The snow had no package and no marks — how strange!",
            "Clues added up: heavy, ugly, and it drinks milk.",
            "Felix looked UP — the present had leaped into the tree!",
        ],
        "suspects": {
            "Rosamond": ["wouldn't say what the present was", "bought six milk cartons for only four cats"],
            "Annie": ["said the gift was \"the most beautiful ever\"", "her dog Fang scared Sludge into a big leap"],
            "Claude": ["saw Rosamond buy an ugly birthday card", "was sitting inside a snow castle"],
            "Sludge": ["leaped super high over a snow pile", "his big leap gave Felix the idea"],
        },
        "evidence_real": ["🪶 Lighter Sled (heavy)", "🥛 Two Extra Milks", "🐾 Sludge's Leap"],
        "evidence_distractor": "⛄ Snow Dog (silly prop — just a decoration Felix built)",
        "deduction": "Sled felt lighter = present is heavy; 6 milks but only 4 cats = something big drinks milk; Sludge's leap = it can leap. Heavy + milk-drinker + leaper = a giant monster cat that jumped off the sled into the tree (so no marks in the snow).",
        "culprit": "Monster Cat (the lost present — Page 4 asks 'what was the present?', not a person)",
    },
    {
        "num": "08",
        "title": "The Fishy Prize",
        "title_cn": "腥味奖品",
        "cover": {
            "Client": "Rosamond",
            "Missing": "The SMARTEST prize — a gold-painted tuna can",
            "Location": "Rosamond's windowsill — vanished mid-stampede",
        },
        "timeline": [
            "Rosamond set her gold \"SMARTEST\" tuna can on the windowsill to dry.",
            "The pets had a wild stampede and messed up the room.",
            "Fang's tail hit the can and knocked it out the open window.",
            "The can dropped into Felix's open bag as he biked past.",
            "At home, Sludge sniffed the grocery bag — the prize was inside!",
        ],
        "suspects": {
            "Fang": ["stood with his right side to the open window", "had gold paint on the right side of his tail"],
            "Anastasia": ["gobbled a big pile of food very fast", "had NO gold paint inside her mouth"],
            "Cats": ["licked the tuna can clean — they love fish", "ran wild and helped wreck the room"],
            "Felix": ["rode past on his new bike during the noise", "left an open grocery bag in his bike basket"],
        },
        "evidence_real": ["🎨 Gold on Fang's Tail", "🚲 Open Grocery Bag", "👃 Sludge's Fishy Sniff"],
        "evidence_distractor": "🐷 Pig Ate It? (ruled-out lead — Anastasia had no gold paint in her mouth)",
        "deduction": "Gold paint on the RIGHT side of Fang's tail means his tail knocked the can OUTWARD, out the window — just as Felix biked past with an open grocery bag in his basket. The can dropped in, and Felix carried the prize home without knowing it (Sludge smelled the fishy can and sniffed it out first).",
        "culprit": "Felix (the detective himself — he carried it home by accident; Page 4 asks 'who took it home without knowing?')",
    },
    {
        "num": "09",
        "title": "The Vanishing Weed",
        "title_cn": "消失的杂草",
        "cover": {
            "Client": "Oliver",
            "Missing": "His adopted weed — small, sick, with a yellow bud",
            "Location": "Oliver's screened-in back porch",
        },
        "timeline": [
            "Oliver bought a sick little weed at Rosamond's Adopt-A-Weed sale.",
            "He potted it and set it in the sun on his porch railing.",
            "Oliver turned away to fetch water — his certificate swung at the pot.",
            "The weed plopped into the open book; the breeze blew it shut.",
            "Felix found the weed pressed flat inside the book at the library.",
        ],
        "suspects": {
            "Oliver": ["bought the sick weed for a nickel", "turned to get water, certificate poking from his pocket"],
            "Rosamond": ["runs an Adopt-A-Weed sale with a song for weeds", "keeps a list of all her weeds in a book"],
            "Fang": ["will eat almost anything", "grabbed a weed and ran off with it in his teeth"],
            "Claude": ["was hunting a lost worm down in the soil", "said it was \"right under our noses\" but unseen"],
        },
        "evidence_real": ["📜 The Certificate", "📖 Open Weed Book", "🌬️ The Breeze"],
        "evidence_distractor": "🐶 Fang Ate It? (ruled-out lead — Fang ran off with a different weed, not this one)",
        "deduction": "Oliver's Certificate of Ownership poked from his back pocket. Turning away to get water, it swung and knocked the sick weed out of its dry pot; the weed plopped into the open weed book right beside it; the breeze then blew the book shut — so the weed was pressed flat inside, right under their noses (just like Claude's hidden worm).",
        "culprit": "The Weed Book (where the weed ended up — Page 4 asks 'where did the weed go?', not a person; Oliver's certificate knocked it in by accident)",
    },
    {
        "num": "10",
        "title": "The Boring Beach Bag",
        "title_cn": "无聊的沙滩包",
        "cover": {
            "Client": "Oliver",
            "Missing": "His boring blue beach bag — clothes, shoes & a seashell inside",
            "Location": "The beach — beside his beach ball",
        },
        "timeline": [
            "Oliver left his boring beach bag and ball on the sand.",
            "He ran to Rosamond's Restaurant for a glass of water.",
            "A kicked beach ball drifted far down the beach.",
            "Oliver followed Esmeralda, spotted a ball, and cried \"stolen!\"",
            "Felix found the bag untouched, where Oliver first left it.",
        ],
        "suspects": {
            "Rosamond": ["sells sandy water and sand sandwiches at her stand", "was too busy with her restaurant to see anything"],
            "Annie": ["ran past with Fang from one end of the beach to the other", "said Fang's on a diet — a bag is not on it"],
            "Esmeralda": ["was hiding from Oliver out in the water", "swam away the moment Oliver showed up"],
            "Oliver": ["left his bag beside his beach ball, then went for water", "follows everyone everywhere — even into the sea"],
        },
        "evidence_real": ["🏖️ Unpressed Sand", "🏐 The Moving Ball", "👣 Following Esmeralda"],
        "evidence_distractor": "🐶 Fang Ate It? (ruled-out lead — Annie says a bag is not on Fang's diet)",
        "deduction": "The sand had no dent where the heavy bag supposedly sat, so the bag never moved from there. What actually moved was the beach BALL — someone kicked it far down the beach. Oliver, busy following Esmeralda away from his spot, saw his ball in its new place and assumed he was back at the start, so he thought the bag was stolen. It had sat untouched the whole time.",
        "culprit": "Oliver — he got mixed up (no one took the bag; the ball moved, and Oliver confused the spot while following Esmeralda). Page 4 asks 'who really lost it?'",
    },
    {
        "num": "11",
        "title": "Down in the Dumps",
        "title_cn": "垃圾场探案",
        "cover": {
            "Client": "Rosamond",
            "Missing": "Her empty white money box, with ROSAMOND printed on one side",
            "Location": "On the grass by her fortune-telling table",
        },
        "timeline": [
            "Rosamond set up her table, sign, box, and tuna cans on the grass.",
            "Rosamond and her cats went inside to get the crystal ball.",
            "Claude carried out the crate and set it down over the box.",
            "Claude tripped over the tuna cans and ran away.",
            "Felix found the box hidden under the crate Rosamond sat on.",
        ],
        "suspects": {
            "Rosamond": ["reads the future with a crystal ball for three cents", "has four cats who love tuna fish"],
            "Annie": ["showed up with Fang, who might bite Sludge someday", "wanted Felix to read her future too"],
            "Finley": ["owns a rat that sleeps in a box until it chews it up", "had a chewed box marked RAT HOUSE beside his rat"],
            "Claude": ["was tired of being Rosamond's moving man and wanted to run off", "carried the heavy crate against his stomach, looking up"],
        },
        "evidence_real": ["📦 Box < Crate", "🔄 Upside-Down Crate", "👀 Never Looked Down"],
        "evidence_distractor": "🐀 RAT HOUSE Box? (ruled-out lead — Felix climbed the dump for it, but it was Finley's rat's chewed box, not Rosamond's)",
        "deduction": "Rosamond's box was smaller than her crate. All the supermarket crates had open tops, so her crate did too — but she kept it upside down to sit on, putting the open top at the bottom. Claude carried the heavy crate against his stomach while looking UP for Rosamond, so he never looked down. He set the open-bottom crate right over the little box and hid it without knowing. Rosamond was sitting on her own box the whole time.",
        "culprit": "Claude — by accident (no one stole the box; Claude set his upside-down, open-bottom crate right over it without looking down, and Rosamond sat on top the whole time). THE TWIST!",
    },
    {
        "num": "12",
        "title": "The Halloween Hunt",
        "title_cn": "万圣节大搜寻",
        "cover": {
            "Client": "Rosamond",
            "Missing": "Little Hex, her small black cat who is scared of Halloween",
            "Location": "Last seen following Rosamond into Esmeralda's house",
        },
        "timeline": [
            "Little Hex followed Rosamond and Annie out trick-or-treating.",
            "At Esmeralda's house, everyone helped her into a gorilla costume.",
            "Little Hex crawled into Rosamond's basket and ate the treats.",
            "Felix searched the spooky old haunted house — no Little Hex.",
            "Felix lifted the red cloth — Little Hex was asleep in the basket!",
        ],
        "suspects": {
            "Esmeralda": ["asked everyone in to help her into her gorilla costume", "kept her door open while the cat slipped away"],
            "Annie": ["dressed up her dog Fang as a grandmother with big teeth", "her basket still had room for more treats"],
            "Claude": ["gave out cookies to the trick-or-treaters", "never went near Esmeralda's house that night"],
            "Rosamond": ["said her basket was too heavy to hold any more treats", "was the last one Little Hex followed before he vanished"],
        },
        "evidence_real": ["⚖️ Too-Heavy Basket", "🙈 Hides When Scared", "🧺 Covered Basket"],
        "evidence_distractor": "👻 Haunted House? (ruled-out lead — Felix searched it, but only the other three Hex cats were there, not Little Hex)",
        "deduction": "Rosamond and Annie trick-or-treated at all the same houses, yet Rosamond's basket was much heavier than Annie's — something extra was inside it. Scared animals hide under things (Sludge hid under a chair, then under a sheet). Little Hex, scared of Halloween, crawled into Rosamond's covered basket at Esmeralda's house, ate most of the treats, and hid under the cloth. Rosamond carried him around all night without knowing.",
        "culprit": "Rosamond — by accident (no one took the cat; Little Hex hid in her covered basket and ate the treats, making it heavier than Annie's; she carried him around all night unaware). Page 4 asks 'who had him the whole time?' THE TWIST!",
    },
    {
        "num": "13",
        "title": "The Musical Note",
        "title_cn": "音乐谜条",
        "cover": {
            "Client": "Pip",
            "Mystery": "A strange riddle-note from his piano teacher — what must Pip do at 4 o'clock?",
            "Location": "Rosamond's garage piano studio",
        },
        "timeline": [
            "Pip found a strange riddle-note instead of his piano lesson.",
            "Pip brought the riddle to Felix to crack before four o'clock.",
            "At the band concert, Sludge's \"dance steps\" gave Felix an idea.",
            "Felix read the note as piano keys: A note → middle C → C-sharp.",
            "\"C-sharp\" means \"see sharp\" — Pip is off to get a haircut!",
        ],
        "suspects": {
            "Pip": ["found a mystery note where his piano lesson should be", "his long hair covers half his face — hard to see"],
            "Rosamond": ["wrote the riddle-note and teaches piano in her garage", "gives gold stars for a good music lesson"],
            "Annie": ["was practicing the old piano with her dog Fang", "is also waiting for her late piano lesson"],
            "Sludge": ["did fake \"dance steps\" dodging a bee at the concert", "let Pip trip over him to give Felix a clue"],
        },
        "evidence_real": ["🅰️ Underlined \"A note\"", "⬅️ Left to Middle C", "⬆️ Step Up = C-sharp"],
        "evidence_distractor": "🎭 The Garage Stage? (ruled-out lead — Felix first guessed step-up-to-the-stage / dance lessons, twice wrong)",
        "deduction": "The note's 'steps' were not dance steps — they were steps on the piano. Rosamond underlined 'A note', so you start at the A key. Step left until you reach the middle = middle C. Step up to the black key = C-sharp. 'C-sharp' sounds just like 'see sharp'. Pip's long hair covers his eyes, so at four o'clock his mother is taking him for a haircut — so he can see sharp.",
        "culprit": "A Haircut (the riddle's answer — no villain; C-sharp = 'see sharp', so Pip goes for a haircut at 4 o'clock). Page 4 asks 'what does the riddle mean?', not who did it.",
    },
    {
        "num": "14",
        "title": "The Stolen Base",
        "title_cn": "失踪的二垒",
        "cover": {
            "Client": "Oliver",
            "Missing": "His good-luck octopus — a purple plastic toy used as second base!",
            "Location": "The baseball field & Oliver's messy room",
        },
        "timeline": [
            "At the game, the octopus was second base.",
            "Walking home, Fang snatched a couple of its arms.",
            "At home, Oliver dumped the octopus on his messy bookcase.",
            "The octopus slid off the back of the bookcase and vanished.",
            "Felix remembered Rosamond's stuck mitt — and found the octopus on the phone cord!",
        ],
        "suspects": {
            "Oliver": ["collects eels and is saving up for a real octopus", "calls everyone he knows on his telephone all day"],
            "Rosamond": ["coaches the team and brings the tuna-can first base", "tossed her mitt and it got stuck up in a tree branch"],
            "Annie": ["brings the dog-bone third base to every game", "is very proud her dog didn't eat the base"],
            "Fang": ["snatched some octopus arms from Oliver's pocket", "buried a dirty octopus arm in Annie's yard"],
        },
        "evidence_real": ["🐙 Long Curling Arms", "📚 Fell Off the Back", "☎️ Cord Down the Back"],
        "evidence_distractor": "⚾ The Babe Ruth Card? (ruled-out lead — Felix found this lost baseball card behind the bookcase, but it wasn't the octopus)",
        "deduction": "Oliver's octopus has long, curling arms that grab onto things. When it slid off the back of his bookcase, it should have landed on the floor — but it never got there. The telephone cord runs down the back of the bookcase, and the octopus's arms caught on it, just like Rosamond's mitt caught on a tree branch instead of falling. The base was never really stolen (Fang only snatched a couple of loose arms).",
        "culprit": "The Phone Cord (the octopus's location — no villain; it was never stolen, its curling arms snagged the telephone cord behind the bookcase). Page 2 flags Fang as THE SNEAK (he did grab a few arms); Page 4 asks 'where did it go?'",
    },
    {
        "num": "15",
        "title": "The Pillowcase",
        "title_cn": "失踪的枕套",
        "cover": {
            "Client": "Rosamond",
            "Missing": "Big Hex the cat's raggy pillowcase — slashed, shrunken, with holes at one end!",
            "Location": "The middle of the night — Rosamond's, Annie's & the diner",
        },
        "timeline": [
            "Rosamond washed the cats' things and hung them out to dry.",
            "Fang growled, so Annie grabbed her things and rushed home.",
            "At 2 a.m. Rosamond woke Felix — the pillowcase was missing!",
            "Felix searched the houses and the diner — no pillowcase.",
            "Felix realized the \"laundry bag\" Annie took was the pillowcase!",
        ],
        "suspects": {
            "Rosamond": ["called Felix at two in the morning about the pillowcase", "washed all four cats' things and hung them out to dry"],
            "Annie": ["left Rosamond's house in a big hurry when Fang growled", "took Little Hex's nightshirt home by mistake too"],
            "Fang": ["growled and showed all his teeth at Rosamond", "wears pajamas, a sweater, and a neck bandanna"],
            "Diner Man": ["wipes his counter with a small shredded rag full of holes", "saves leftover bones for all the town's dogs"],
        },
        "evidence_real": ["🧺 Open End + Holes", "🏃 Annie's Big Hurry", "❓ No Bag Is Missing"],
        "evidence_distractor": "🧽 The Diner Rag? (ruled-out lead — Felix grabbed the diner man's shredded rag, but it wouldn't open up; it was just a rag)",
        "deduction": "Big Hex's raggy pillowcase was open at one end with holes around it — so it looked just like a laundry bag. Annie was in a big hurry when Fang growled, so she grabbed it and stuffed Fang's clothes inside, thinking it was a bag. But Rosamond was never missing a laundry bag, so Annie never took a real one. The 'laundry bag' she carried home WAS the pillowcase. (Just like Sludge cared about the bone, not the bag it was in.)",
        "culprit": "Annie's \"Laundry Bag\" (the pillowcase itself — no villain; Annie took it by mistake, thinking the open, holey pillowcase was a laundry bag, and stuffed Fang's clothes in it). Page 2 flags Annie as THE MIX-UP; Page 4 asks 'where is the pillowcase?'",
    },
    {
        "num": "16",
        "title": "The Mushy Valentine",
        "title_cn": "肉麻的情人节卡",
        "cover": {
            "Client": "Sludge",
            "Mystery": "A mushy valentine appeared on Sludge's doghouse, signed \"ABH\" — who sent it?",
            "Location": "Sludge's doghouse & Annie's house",
        },
        "timeline": [
            "Annie made a valentine for Harry and signed just \"A\".",
            "Hungry Fang lured Annie away; her valentine vanished.",
            "Harry changed it — adding words and \"B\" + \"H\" to the \"A\".",
            "Harry taped the new valentine onto Sludge's doghouse.",
            "Felix proved ABH = Annie's Brother Harry — both cases were one!",
        ],
        "suspects": {
            "Annie": ["made a paper-heart valentine and signed just the letter \"A\"", "just loves to give valentines to everyone"],
            "Rosamond": ["made a matching \"valentwin\" heart that looked the same", "turned her cats' leftover liver into a valentine"],
            "Harry": ["does not like getting valentines at all", "disappeared the moment the valentine disappeared"],
            "Fang": ["looked hungry and lured Annie out to the kitchen", "will never be anybody's valentine"],
        },
        "evidence_real": ["💌 Matching Valentine", "👃 Sniffed Annie's Desk", "🅰️ Only \"A\" Was Signed"],
        "evidence_distractor": "🐱 \"A Big Hex\"? (ruled-out lead — Felix first guessed ABH meant 'A Big Hex', Rosamond's cat, but ruled it out)",
        "deduction": "Sludge's valentine looked exactly like the matching 'valentwin' hearts Annie and Rosamond made. Sludge kept sniffing Annie's desk — where her valentine had been — because his valentine was once there. Annie signed only the letter 'A' (meaning to write ANNIE). Her brother Harry added a 'B' and an 'H', making ABH = Annie's Brother Harry. Harry didn't want the valentine but couldn't waste something good (like Rosamond's leftover liver, changed into a cat valentine), so he changed it into a valentine for Sludge. Annie's missing valentine and Sludge's valentine were the same one.",
        "culprit": "Harry (Annie's Brother Harry = ABH). The two cases — Annie's missing valentine and Sludge's mystery valentine — are the same case. Page 2 flags Harry as ABH; Page 4 is a culprit-drag (Harry).",
    },
    {
        "num": "17",
        "title": "The Tardy Tortoise",
        "title_cn": "迟到的乌龟",
        "cover": {
            "Client": "Felix (a lost tortoise!)",
            "Mystery": "A slow green tortoise is eating Felix's garden flowers — where does he live?",
            "Location": "Felix's garden & the neighborhood yards",
        },
        "timeline": [
            "A green tortoise was eating the flowers in Felix's garden.",
            "Felix asked Rosamond, Claude, and the vet — no luck.",
            "Felix spotted u-shaped bite marks — like Sludge's crumb trail.",
            "They followed the bite trail backward through yard after yard.",
            "They reached the \"Beware of the Tortoise\" yard — Speedy was home!",
        ],
        "suspects": {
            "Rosamond": ["owns four cats but no tortoise", "was making crumbly tuna-fish cupcakes"],
            "Claude": ["lost something — but it turned out to be his sock, not a pet", "wants a matching size eleven-and-a-half sock"],
            "Annie": ["was at the vet because Fang had a sore tooth", "says the vet won't treat reptiles like tortoises"],
            "Speedy": ["left a trail of u-shaped bite marks on the flowers", "crawls very, very slowly and keeps his secrets"],
        },
        "evidence_real": ["🐢 U-Shaped Bite Marks", "🐾 Follow the Trail Back", "🪧 \"Beware of the Tortoise\" Sign"],
        "evidence_distractor": "🧦 Claude's Lost Sock? (ruled-out lead — Felix found Claude's sock on the trail, but it belonged to a different case)",
        "deduction": "The tortoise ate his way to Felix's garden, leaving little u-shaped bite marks on flower after flower. That was a trail, just like the cupcake-crumb trail Sludge had sniffed. So Felix followed the bite marks backward, yard by yard, until he reached a yard with a fence sign reading 'BEWARE OF THE TORTOISE' — the home of Speedy, the escape-artist tortoise.",
        "culprit": "The \"Beware of the Tortoise\" yard (the tortoise's home — no villain; Speedy escaped under the fence and ate his way to Felix's garden). Page 2 flags Speedy as THE TRAIL-MAKER; Page 4 asks 'where is Speedy's home?'",
    },
    {
        "num": "18",
        "title": "The Crunchy Christmas",
        "title_cn": "嘎吱响的圣诞节",
        "cover": {
            "Client": "Annie (for Fang)",
            "Missing": "Fang's yearly Christmas card from his mother — it never arrived!",
            "Location": "Annie's house & the snowy backyard",
        },
        "timeline": [
            "Fang's yearly card from his mother never showed up.",
            "Fang chased the mailman, so the mail fell on the ground.",
            "Fang smelled a bone in a thick envelope and buried it in the yard.",
            "Felix searched Annie's catalogs and Rosamond's cat-alog — no card.",
            "Felix dug up a snow hump — the bone with Mom's \"Eat!\" card!",
        ],
        "suspects": {
            "Annie": ["collects huge stacks of holiday catalogs", "reads Fang his card on her lap each Christmas"],
            "Fang": ["chases the mailman and buries things in the backyard", "wears jingle bells and an elf hat at Christmas"],
            "Rosamond": ["decorates a cat tree with painted tuna-fish cans", "borrowed Annie's cat catalog last week"],
            "Mailman": ["drops the mail and flees when Fang charges out", "sometimes leaves the mail on the ground by the box"],
        },
        "evidence_real": ["🦴 \"Eat Bones!\" Notes", "✉️ Thick Heavy Envelope", "🐕 Dogs Bury Bones"],
        "evidence_distractor": "💉 The Vet Postcard? (ruled-out lead — addressed to Fang, fell out of a catalog, but it was just a shot reminder, not the missing card)",
        "deduction": "Every year Mrs. Fang's cards nagged Fang harder to eat his bones. This year's message was the strongest: she mailed a real bone in a thick envelope. When Fang chased the mailman, the mail fell on the ground. Fang sniffed the envelope, smelled the bone inside, and did what dogs do — buried it in the backyard. Then the snow covered the hump. Felix had to 'think like a dog' to realize Fang was hiding his own card.",
        "culprit": "Fang (no real thief; he buried his own Christmas card because it contained a real bone). Page 2 flags Fang as THE BURIER; Page 4 is a culprit-drag (Fang).",
    },
    {
        "num": "19",
        "title": "Saves the King of Sweden",
        "title_cn": "拯救瑞典国王",
        "cover": {
            "Client": "Rosamond (visiting Scandinavia)",
            "Missing": "A troll figurine — a present she bought and hid",
            "Location": "Scandinavia... but Felix solves it from home",
        },
        "timeline": [
            "Rosamond's postcard arrived — she lost something smart with a very long nose.",
            "Felix went to the library and learned trolls live in dark places and love berries.",
            "A palace photo came; Felix first thought Rosamond held the lost troll.",
            "Felix looked closer and saw Rosamond was holding one of her cats, not the troll.",
            "Felix remembered trolls like dark places — the troll was in Rosamond's right hiking boot!",
        ],
        "suspects": {
            "Rosamond": ["visiting Scandinavia and bought a troll present", "wrote that the lost thing was smart with a very long nose"],
            "Annie": ["received Rosamond's picture postcard", "showed Felix the photo from the palace"],
            "Esmeralda": ["knew Rosamond was traveling in Scandinavia", "said trolls are very good at hiding in dark places"],
            "Sludge": ["buried his bone in a special hiding place", "forgot where he hid the bone until Felix helped him find it"],
        },
        "evidence_real": ["📮 Rosamond's Postcard", "🌑 Trolls Love Dark Places", "🦴 Sludge Buries His Bone"],
        "evidence_distractor": "🐱 Palace Photo (ruled-out lead — Rosamond is holding a cat, not the troll)",
        "deduction": "Rosamond's postcard said the lost thing was smart with a very long nose. Trolls have long noses and love dark hiding places. Sludge buried his bone in a special place and forgot where it was — just like Rosamond did. So the troll must be in a dark hiding place Rosamond forgot: her right hiking boot.",
        "culprit": "Rosamond's Right Hiking Boot (Rosamond hid the troll there because trolls like dark places, then forgot). Page 2 flags Rosamond as THE HIDER; Page 4 is a location-drag (the boot).",
    },
    {
        "num": "20",
        "title": "The Monster Mess",
        "title_cn": "怪物饼干之乱",
        "cover": {
            "Client": "Felix's Mother",
            "Missing": "The Monster Cookie recipe (Strawberry Draculas, Chocolate Frankensteins, Cinnamon Werewolves)",
            "Location": "Felix's kitchen and neighborhood",
        },
        "timeline": [
            "Felix's mother wanted to bake Monster Cookies but her recipe was gone.",
            "Felix and Sludge searched the kitchen, sniffing and looking everywhere.",
            "They found crinkly, fishy, fangy clues all over the house and street.",
            "Oliver said he saw Mom at the grocery store three days earlier.",
            "Felix found the recipe — Mom had scribbled it on the back of another recipe!",
        ],
        "suspects": {
            "Mother": ["wanted to bake Strawberry Dracula, Chocolate Frankenstein, and Cinnamon Werewolf cookies", "had looked everywhere but could not find the recipe"],
            "Oliver": ["follows everyone everywhere", "saw Mom at the grocery store three days earlier"],
            "Annie": ["can stay home alone and keep secrets", "gave Mom a brand-new monster cookie cutter"],
            "Sludge": ["sniffed long and short clues all around the house", "found a crinkly, wrinkly paper under the couch"],
        },
        "evidence_real": ["🍪 Special Monster Names", "📝 Mom's Recipe Cards", "✏️ Mom Scribbles on Backs"],
        "evidence_distractor": "🐟 Fishy Clue (ruled-out lead — just the fish store, not the recipe)",
        "deduction": "Monster Cookies have special names, so Mom needed the recipe. She keeps her recipes on cards. And Mom likes to scribble notes on the back of things. So the recipe was on the back of another recipe card all along.",
        "culprit": "The Back of Another Recipe (Mom had scribbled the Monster Cookie recipe there and forgotten). Page 2 flags Mother as THE FORGETFUL WRITER; Page 4 is a location-drag (the back of a recipe card).",
    },
    {
        "num": "21",
        "title": "San Francisco Detective",
        "title_cn": "旧金山侦探",
        "cover": {
            "Client": "Duncan (Olivia Sharp's chauffeur)",
            "Missing": "Duncan's favorite joke book — 'Joke Stew'",
            "Location": "San Francisco — pancake house, Golden Gate Bridge, Lombard Street, and more",
        },
        "timeline": [
            "Felix and Sludge arrived in San Francisco to visit cousin Olivia Sharp.",
            "Olivia was too busy, so Duncan asked Felix to find his lost joke book.",
            "Felix and Sludge rode around San Francisco retracing Duncan's steps.",
            "They visited a pancake house, the Golden Gate Bridge, and Lombard Street.",
            "Felix found 'Joke Stew' in the cookbook section at Booksie's Bookstore!",
        ],
        "suspects": {
            "Duncan": ["Olivia Sharp's chauffeur and client", "lost his favorite joke book 'Joke Stew'"],
            "Olivia": ["Felix's cousin and a busy detective", "sent her limousine to greet Felix at the airport"],
            "Booksie": ["owns Booksie's Bookstore in San Francisco", "the store has a cookbook section"],
            "Sludge": ["Felix's loyal dog and detective partner", "sniffed clues all over San Francisco"],
        },
        "evidence_real": ["🚗 Retracing Duncan's Steps", "📚 Booksie's Bookstore", "🙃 Seems Wrong but Right"],
        "evidence_distractor": "🥞 Pancake House (ruled-out lead — they only ate pancakes there)",
        "deduction": "Felix retraced Duncan's steps across San Francisco. The last clues led to Booksie's Bookstore. A joke book seems out of place there, but 'a place that seems wrong but could be right' means the cookbook section — exactly where a joke book would be misplaced.",
        "culprit": "Booksie's Cookbook Section (Duncan's 'Joke Stew' was misplaced among the cookbooks). Page 2 flags Booksie as THE BOOKSPOT; Page 4 is a location-drag (the cookbook section).",
    },
    {
        "num": "22",
        "title": "The Big Sniff",
        "title_cn": "大鼻子嗅探",
        "cover": {
            "Client": "Felix himself (Sludge is missing!)",
            "Missing": "Sludge, lost inside Weinman Brothers department store",
            "Location": "Weinman Brothers department store, on a rainy day",
        },
        "timeline": [
            "Felix went into Weinman Brothers to buy a gift for Sludge.",
            "Sludge was left outside because dogs are not allowed.",
            "After a long line, Felix went outside and found Sludge gone.",
            "Salespeople reported a wet, slippery, sloppy dog running around inside.",
            "Felix used Sludge's own sniffing skills to find him in the underwear department.",
        ],
        "suspects": {
            "Felix": ["went into Weinman Brothers to buy Sludge a gift", "left Sludge outside because no dogs are allowed"],
            "Rosamond": ["was shopping in the department store", "was looking at clothes for her cats"],
            "Salespeople": ["work at Weinman Brothers", "reported a wet, slippery, sloppy dog running around"],
            "Sludge": ["Felix's loyal dog and detective partner", "has a very strong sense of smell"],
        },
        "evidence_real": ["👃 Sludge's Strong Nose", "🌧️ Rainy, Wet Day", "🙃 Seems Wrong but Right"],
        "evidence_distractor": "🛏️ Mattress Department (ruled-out lead — Sludge was not there)",
        "deduction": "Sludge has a very strong nose. It was raining, so he was wet and slipped inside when someone opened the door. Felix followed Sludge's sniffing trail through the store. The trail led to the underwear department — a place that seems wrong for a dog but was just right for hiding!",
        "culprit": "The Underwear Department (Sludge slipped inside and hid there). Page 2 flags Sludge as THE SNIFFER; Page 4 is a location-drag (the underwear department).",
    },
    {
        "num": "23",
        "title": "The Owl Express",
        "title_cn": "猫头鹰快车",
        "cover": {
            "Client": "Olivia Sharp",
            "Missing": "Hoot the owl — vanished from his cage on the train",
            "Location": "The Owl Express train to San Francisco",
        },
        "timeline": [
            "Olivia Sharp hired Felix to guard her pet owl Hoot on the Owl Express train.",
            "A threatening note said it would be a happy day when Hoot flies away.",
            "Felix fell asleep, and when he woke up, Hoot had vanished from the cage.",
            "Felix questioned the sneezing lady, waiter, musician, and owl man.",
            "Felix found Hoot on the roof of the train — Hoot had flown out of his cage!",
        ],
        "suspects": {
            "Olivia": ["hired Felix to guard Hoot on the Owl Express", "received the note and was worried about Hoot"],
            "Musician": ["wrote a threatening note about Hoot flying away", "played music that made Hoot's feathers flutter"],
            "Waiter": ["served food and drinks in the dining car", "saw Hoot's cage while carrying a tray"],
            "OwlMan": ["knows a lot about owls and their habits", "was watching Hoot's cage very closely"],
        },
        "evidence_real": ["🦉 Hoot Has Wings", "🚂 Open Train Roof", "📝 Note: 'Fly Away'"],
        "evidence_distractor": "🎵 The Musician (ruled-out lead — wrote the note but did not take Hoot)",
        "deduction": "Hoot is an owl with wings, so he can fly. The train roof is open to the sky. The note said it would be a happy day when Hoot flies away. So Hoot did not need to be stolen — he flew out of his cage and landed on the roof of the train.",
        "culprit": "The Train Roof (Hoot flew out of his cage and landed there). Page 2 flags OwlMan as THE WATCHER; Page 4 is a location-drag (the train roof).",
    },
    {
        "num": "24",
        "title": "Talks Turkey",
        "title_cn": "火鸡追踪",
        "cover": {
            "Client": "Claude (and Sludge took the case!)",
            "Missing": "A big white turkey — found in the woods, then lost again",
            "Location": "Deering Woods → 58 Kenwood Street",
        },
        "timeline": [
            "Felix heard on the radio that a giant turkey caused panic in a parking lot.",
            "Claude found a big white turkey in the woods, then lost it.",
            "Sludge sniffed Claude's white feather and took the case.",
            "On TV, the famous turkey had dark green and red feathers — two turkeys!",
            "Sludge tracked the white turkey home to 58 Kenwood Street.",
        ],
        "suspects": {
            "Claude": ["found a big white turkey in Deering Woods", "dropped popcorn, so the turkey followed him"],
            "Olivia": ["Felix's cousin, a bird-loving detective from San Francisco", "came to catch the famous parking-lot turkey"],
            "Rosamond": ["runs a Rent-A-Pet stand for a nickel an hour", "had a big white turkey from Kenwood Street on her list"],
            "Sludge": ["sniffed the white feather Claude found", "waited by the tracks so he would not mess them up"],
        },
        "evidence_real": ["🪶 White Feather", "📺 Colorful TV Turkey", "👃 Sludge's Nose"],
        "evidence_distractor": "🍗 Thanksgiving Dinner? (ruled-out lead — it's summer, Thanksgiving is months away)",
        "deduction": "Everyone thought ONE turkey was missing. But look for what doesn't fit — the feathers! Claude's feather was WHITE, but the famous turkey on TV had DARK GREEN and RED feathers. Two different feathers = TWO turkeys. Sludge followed the white feather's smell and tracked that turkey home to 58 Kenwood Street; Olivia caught the famous colorful turkey with turkey food.",
        "culprit": "Two Turkeys (the twist — there were two different turkeys all along; Claude's was white, the famous one colorful). Page 2 flags Sludge as THE TRACKER; Page 4 asks 'how many turkeys were missing?'",
    },
    {
        "num": "25",
        "title": "The Hungry Book Club",
        "title_cn": "饥饿读书会",
        "cover": {
            "Client": "Rosamond (president of the Ready Readers)",
            "Mystery": "An 'Evil Page Monster' tore a cookbook page — torn, ripped, ruined!",
            "Location": "Rosamond's book club — the kitchen table",
        },
        "timeline": [
            "Rosamond started a book club, but a cookbook page was torn, ripped, ruined!",
            "Felix and Sludge went undercover as new book-club members.",
            "Rosamond blamed Fang and called him the Evil Page Monster.",
            "Felix saw the table was too high — Fang could not reach the cookbook.",
            "Felix proved a small, leaping cat ripped the page for the tuna bits.",
        ],
        "suspects": {
            "Fang": ["Annie's big dog — Rosamond blamed him first", "too short for the high table, and hates tuna pie"],
            "LittleHex": ["Rosamond's smallest cat, and cats love tuna fish", "can leap high and hide in tiny spaces unseen"],
            "Oliver": ["follows everyone everywhere in the book club", "said Harvard Hedgehog is a very boring book"],
            "Esmeralda": ["was reading aloud when she found a page missing", "first thought Harvard just needed a watch"],
        },
        "evidence_real": ["📏 Table Too High", "🐟 Tuna Bits Gone", "🐱 Small Leaping Cat"],
        "evidence_distractor": "🐶 Fang the Monster? (ruled-out lead — the table's too high for him and he hates tuna pie)",
        "deduction": "Rosamond blamed Fang, but the kitchen table is too HIGH — Fang jumped and stretched but couldn't reach the cookbook, and he doesn't even like tuna fish pie. The torn page had a tuna recipe, and all the tuna bits were scratched away. Only a small cat that can LEAP high, HIDE in tiny spaces, and sneak in unseen fits — Little Hex tore the page to get every last bite of tuna.",
        "culprit": "Little Hex (Rosamond's smallest cat; he leaped onto the high table and tore the page for the tuna bits — Fang was wrongly blamed). Page 2 flags Little Hex as THE PAGE MONSTER; Page 4 is a culprit-drag (Little Hex).",
    },
]

BROWN = RGBColor(0x4A, 0x35, 0x28)
RED = RGBColor(0xB8, 0x3A, 0x2A)


def add_heading(doc, text, size, color=BROWN, bold=True, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_label_line(doc, label, value, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RED
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    return p


def build():
    doc = Document()

    add_heading(doc, "侦探案件答案验证表", 22)
    add_heading(doc, "Detective Case Files — Answer Key", 14, color=BROWN, bold=False, after=4)
    intro = doc.add_paragraph()
    intro.add_run(
        "每个案件按 5 个页面列出正确答案，方便逐页验证网页功能。"
        "（每新建一个案件/页面，本表会同步增加。）"
    ).italic = True
    intro.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    for c in CASES:
        add_heading(doc, f"Case {c['num']} — {c['title']}  {c['title_cn']}", 17, before=8)

        # Page 0 — Cover (context, not a puzzle)
        add_heading(doc, "Page 0 · Cover（封面信息）", 12, color=BROWN)
        for k, v in c["cover"].items():
            add_label_line(doc, f"{k}:", v, indent=12)

        # Page 1 — Timeline
        add_heading(doc, "Page 1 · Timeline 时间线（正确顺序 1→5）", 12, color=BROWN)
        for i, ev in enumerate(c["timeline"], 1):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Pt(20)
            p.add_run(ev).font.size = Pt(11)

        # Page 2 — Suspect Files
        add_heading(doc, "Page 2 · Suspect Files 嫌疑人文件（线索 → 文件夹）", 12, color=BROWN)
        for name, clues in c["suspects"].items():
            tag = "  ← 真凶 CULPRIT" if name == c["culprit"] else ""
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(name + tag)
            r.bold = True
            r.font.size = Pt(11)
            if name == c["culprit"]:
                r.font.color.rgb = RED
            for clue in clues:
                pc = doc.add_paragraph(style="List Bullet")
                pc.paragraph_format.left_indent = Pt(30)
                pc.paragraph_format.space_after = Pt(0)
                pc.add_run(clue).font.size = Pt(10.5)

        # Page 3 — Evidence Board
        add_heading(doc, "Page 3 · Evidence Board 证据板", 12, color=BROWN)
        add_label_line(doc, "拖上板 (real):", "  ·  ".join(c["evidence_real"]), indent=12)
        add_label_line(doc, "拒绝/弹回 (distractor):", c["evidence_distractor"], indent=12)
        add_label_line(doc, "推理 (deduction):", c["deduction"], indent=12)

        # Page 4 — Solve
        add_heading(doc, "Page 4 · Solve the Case 指认真凶", 12, color=BROWN)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        r = p.add_run("Culprit:  " + c["culprit"])
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RED

        doc.add_paragraph()

    doc.save(OUT)
    print("wrote", OUT, "with", len(CASES), "cases")


if __name__ == "__main__":
    build()
